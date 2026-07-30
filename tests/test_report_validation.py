from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from archiv.hashing import sha256_file
from archiv.report_contracts import ReportManifest
from archiv.reports import generate_report, validate_report
from report_support import prepare_report_archive


def _rewrite_manifest_hash(path: Path, docx_path: Path) -> None:
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["docx_sha256"] = sha256_file(docx_path)
    path.write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def test_missing_and_malformed_outputs_never_validate(tmp_path: Path) -> None:
    missing = validate_report(tmp_path / "missing.docx", tmp_path / "missing.json")
    assert missing.valid is False
    assert "required DOCX output is missing" in missing.errors

    malformed = tmp_path / "malformed.docx"
    malformed.write_bytes(b"not a DOCX package")
    manifest = tmp_path / "manifest.json"
    manifest.write_text("{}", encoding="utf-8")
    invalid = validate_report(malformed, manifest)
    assert invalid.valid is False


def test_uncited_report_is_rejected(
    ingestion_corpus: Path,
    tmp_path: Path,
) -> None:
    home = tmp_path / "archiv-home"
    prepare_report_archive(ingestion_corpus, home)
    output = tmp_path / "report.docx"
    generated = generate_report("MARKER", output, home=home, max_sources=3, render=False)
    manifest_path = Path(generated.manifest_path)
    manifest = ReportManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))

    document = Document(str(output))
    target = f"[{manifest.sources[0].number}]"
    for paragraph in document.paragraphs:
        if target in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(target, "")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(target, "")
    document.save(str(output))
    _rewrite_manifest_hash(manifest_path, output)

    validation = validate_report(output, manifest_path, home=home, render=False)
    assert validation.valid is False
    assert f"inline citation {target} missing from its finding" in validation.errors
