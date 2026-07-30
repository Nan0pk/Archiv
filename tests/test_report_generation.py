from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from report_support import REPORT_FIXTURES, prepare_report_archive

from archiv.hashing import sha256_file
from archiv.report_contracts import ReportManifest, ReportStatus
from archiv.reports import generate_report


def _document_text(path: Path) -> str:
    document = Document(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells)
    return "\n".join(blocks)


def test_generates_structurally_valid_cited_docx_without_source_changes(
    ingestion_corpus: Path,
    tmp_path: Path,
) -> None:
    home = tmp_path / "archiv-home"
    originals = prepare_report_archive(ingestion_corpus, home)
    output = tmp_path / "evidence-report.docx"

    result = generate_report(
        "MARKER",
        output,
        home=home,
        max_sources=len(REPORT_FIXTURES),
        render=False,
    )

    assert result.status is ReportStatus.SUCCEEDED
    assert result.validation.valid is True
    assert output.is_file()
    manifest = ReportManifest.model_validate_json(
        Path(result.manifest_path).read_text(encoding="utf-8")
    )
    assert len(manifest.sources) == len(REPORT_FIXTURES)
    assert manifest.docx_sha256 == sha256_file(output)

    text = _document_text(output)
    for section in manifest.required_sections:
        assert section in text
    for source in manifest.sources:
        assert f"[{source.number}]" in text
        assert source.excerpt in text
        assert source.citation.segment_id in text

    validation_payload = json.loads(Path(result.validation_path).read_text(encoding="utf-8"))
    assert validation_payload["valid"] is True
    assert {path: sha256_file(path) for path in originals} == originals
