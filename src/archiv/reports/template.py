"""Default Archiv report template and document styling."""

from __future__ import annotations

from typing import cast

from docx import Document as open_document
from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle


def new_report_document(*, title: str, report_id: str) -> Document:
    """Create the default report template with stable styles and page geometry."""

    document = open_document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = cast(ParagraphStyle, styles["Normal"])
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Title", 22), ("Heading 1", 15), ("Heading 2", 12)):
        style = cast(ParagraphStyle, styles[name])
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.bold = True

    properties = document.core_properties
    properties.title = title
    properties.subject = "Evidence-backed Archiv report"
    properties.author = "Archiv"
    properties.keywords = f"Archiv report {report_id}"

    header = section.header.paragraphs[0]
    header.text = "Archiv Evidence Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.text = f"Report ID: {report_id}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Generated from validated Archiv citations")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    document.add_section(WD_SECTION.CONTINUOUS)
    return document
