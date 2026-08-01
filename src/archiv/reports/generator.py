"""Generate cited DOCX reports from validated Archiv search results."""

from __future__ import annotations

import json
import os
from datetime import UTC, datetime
from pathlib import Path
from uuid import uuid4

from docx.document import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from archiv.contracts import SearchResult
from archiv.grounding_contracts import GroundedModelResponse
from archiv.hashing import sha256_file
from archiv.report_contracts import (
    REQUIRED_REPORT_SECTIONS,
    ReportGenerationResult,
    ReportManifest,
    ReportSource,
    ReportStatus,
)
from archiv.reports.formatting import format_locator
from archiv.reports.template import new_report_document
from archiv.reports.validation import validate_report, write_validation
from archiv.search import read_source_excerpt, search_documents, validate_citation


def _write_manifest(path: Path, manifest: ReportManifest) -> None:
    path.write_text(
        json.dumps(manifest.model_dump(mode="json"), indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def _distinct_sources(results: list[SearchResult], *, limit: int) -> list[SearchResult]:
    selected: list[SearchResult] = []
    seen: set[str] = set()
    for result in results:
        digest = result.citation.object_sha256
        if digest in seen:
            continue
        selected.append(result)
        seen.add(digest)
        if len(selected) == limit:
            break
    return selected


def _add_source_overview(document: Document, sources: list[ReportSource]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Citation", "Source", "Location")
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for source in sources:
        row = table.add_row().cells
        row[0].text = f"[{source.number}]"
        row[1].text = source.citation.source_name
        row[2].text = source.locator_text


def _build_document(
    *,
    title: str,
    query: str,
    report_id: str,
    sources: list[ReportSource],
    grounded_response: GroundedModelResponse | None = None,
    model_identity: str = "disabled",
) -> Document:
    document = new_report_document(title=title, report_id=report_id)

    document.add_heading("Executive Summary", level=1)
    summary = document.add_paragraph()
    summary.add_run(
        f'This report presents {len(sources)} validated source finding(s) for the query "{query}". '
    )
    summary.add_run(
        "Every excerpt and source location was validated against Archiv's immutable original and "
        "normalized evidence before document generation."
    )
    _add_source_overview(document, sources)

    document.add_heading("Findings", level=1)

    cit_map = {f"CIT-{source.number}": f"[{source.number}]" for source in sources}

    if grounded_response:
        for g_para in grounded_response.paragraphs:
            p = document.add_paragraph()
            p.add_run(g_para.text)
            for cid in g_para.citation_ids:
                if cid in cit_map:
                    run = p.add_run(f" {cit_map[cid]}")
                    run.bold = True
                    run.font.size = Pt(9)
        for claim in grounded_response.claims:
            p = document.add_paragraph()
            p.add_run(claim.statement)
            for cid in claim.citation_ids:
                if cid in cit_map:
                    run = p.add_run(f" {cit_map[cid]}")
                    run.bold = True
                    run.font.size = Pt(9)

        if grounded_response.insufficient_evidence:
            document.add_heading("Insufficient Evidence & Unresolved Aspects", level=2)
            for item in grounded_response.insufficient_evidence:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(item)

        if grounded_response.contradictions:
            document.add_heading("Contradictions Between Sources", level=2)
            for item in grounded_response.contradictions:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(item)
    else:
        for source in sources:
            document.add_heading(f"Finding {source.number}: {source.citation.source_name}", level=2)
            doc_paragraph = document.add_paragraph()
            doc_paragraph.add_run(source.excerpt)
            citation_run = doc_paragraph.add_run(f" [{source.number}]")
            citation_run.bold = True
            citation_run.font.size = Pt(9)
            locator = document.add_paragraph()
            locator_run = locator.add_run(f"Source location: {source.locator_text}")
            locator_run.italic = True
            locator_run.font.size = Pt(9)

    document.add_page_break()
    document.add_heading("Source Appendix", level=1)
    for source in sources:
        document.add_heading(f"[{source.number}] {source.citation.source_name}", level=2)
        details = (
            ("Location", source.locator_text),
            ("Object SHA-256", source.citation.object_sha256),
            ("Segment ID", source.citation.segment_id),
            ("Normalized SHA-256", source.citation.normalized_sha256),
            ("Text SHA-256", source.citation.text_sha256),
        )
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in details:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].bold = True
        excerpt = document.add_paragraph()
        excerpt.add_run("Validated excerpt: ").bold = True
        excerpt.add_run(source.excerpt)

    document.add_heading("Provenance", level=1)
    provenance = document.add_paragraph()
    provenance.add_run("Report ID: ").bold = True
    provenance.add_run(report_id)
    provenance.add_run("\nQuery: ").bold = True
    provenance.add_run(query)
    provenance.add_run("\nModel identity: ").bold = True
    provenance.add_run(model_identity)
    provenance.add_run("\nGeneration policy: ").bold = True
    provenance.add_run(
        "validated citations only; no source mutation; DOCX success requires "
        "independent package validation"
    )

    for paragraph in document.paragraphs:
        paragraph_style = paragraph.style
        if paragraph_style is not None and (paragraph_style.name or "").startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def generate_report(
    query: str,
    output: Path,
    *,
    title: str = "Archiv Evidence Report",
    home: Path | None = None,
    search_limit: int = 50,
    max_sources: int = 8,
    render: bool = False,
    evidence_dir: Path | None = None,
    grounded_response: GroundedModelResponse | None = None,
    model_identity: str = "disabled",
) -> ReportGenerationResult:
    """Search Archiv and generate a report from independently validated results."""

    results = search_documents(query, home=home, limit=search_limit)
    return generate_report_from_results(
        results,
        output,
        query=query,
        title=title,
        home=home,
        max_sources=max_sources,
        render=render,
        evidence_dir=evidence_dir,
        grounded_response=grounded_response,
        model_identity=model_identity,
    )


def generate_report_from_results(
    results: list[SearchResult],
    output: Path,
    *,
    query: str,
    title: str = "Archiv Evidence Report",
    home: Path | None = None,
    max_sources: int = 8,
    render: bool = False,
    evidence_dir: Path | None = None,
    grounded_response: GroundedModelResponse | None = None,
    model_identity: str = "disabled",
) -> ReportGenerationResult:
    """Generate a report from validated results and return only evidence-backed status."""

    if max_sources < 1 or max_sources > 50:
        raise ValueError("max_sources must be between 1 and 50")
    selected = _distinct_sources(results, limit=max_sources)
    if not selected:
        raise ValueError("report generation requires at least one search result")

    sources: list[ReportSource] = []
    for number, result in enumerate(selected, 1):
        validation = validate_citation(result.citation, home=home)
        if not validation.valid:
            raise ValueError(
                f"search result {number} has an invalid citation: " + "; ".join(validation.errors)
            )
        excerpt = read_source_excerpt(result.citation, home=home)
        if excerpt != result.text:
            raise ValueError(f"search result {number} text does not match source excerpt")
        sources.append(
            ReportSource(
                number=number,
                citation=result.citation,
                excerpt=excerpt,
                locator_text=format_locator(result.citation.locator),
            )
        )

    output = output.expanduser().resolve()
    if output.suffix.lower() != ".docx":
        raise ValueError("report output must use the .docx extension")
    output.parent.mkdir(parents=True, exist_ok=True)
    report_id = uuid4().hex
    manifest_path = output.with_suffix(output.suffix + ".manifest.json")
    validation_path = output.with_suffix(output.suffix + ".validation.json")
    temporary = output.with_name(f".{output.name}.{report_id}.tmp")

    document = _build_document(
        title=title,
        query=query,
        report_id=report_id,
        sources=sources,
        grounded_response=grounded_response,
        model_identity=model_identity,
    )
    document.save(str(temporary))
    os.replace(temporary, output)

    manifest = ReportManifest(
        report_id=report_id,
        title=title,
        query=query,
        generated_at=datetime.now(UTC).isoformat(),
        docx_path=str(output),
        docx_sha256=sha256_file(output),
        required_sections=list(REQUIRED_REPORT_SECTIONS),
        sources=sources,
    )
    _write_manifest(manifest_path, manifest)
    report_validation = validate_report(
        output,
        manifest_path,
        home=home,
        render=render,
        evidence_dir=evidence_dir,
    )
    write_validation(validation_path, report_validation)
    status = ReportStatus.SUCCEEDED if report_validation.valid else ReportStatus.FAILED
    return ReportGenerationResult(
        status=status,
        report_id=report_id,
        docx_path=str(output),
        manifest_path=str(manifest_path),
        validation_path=str(validation_path),
        validation=report_validation,
    )
