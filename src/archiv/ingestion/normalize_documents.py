"""Text, PDF, and DOCX normalization."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from archiv.contracts import NormalizedDocument, NormalizedSegment
from archiv.ingestion.normalize_pdf import normalize_pdf

__all__ = ["normalize_docx", "normalize_pdf", "normalize_text"]


def normalize_text(
    path: Path,
    digest: str,
    *,
    source_name: str,
    kind: str,
    media_type: str,
) -> NormalizedDocument:
    text = path.read_text(encoding="utf-8")
    segments = [
        NormalizedSegment(locator={"line": index}, text=line)
        for index, line in enumerate(text.splitlines(), 1)
        if line
    ]
    return NormalizedDocument(
        object_sha256=digest,
        media_type=media_type,
        kind=kind,
        source_name=source_name,
        segments=segments,
        metadata={"encoding": "utf-8"},
    )


def normalize_docx(
    path: Path,
    digest: str,
    *,
    source_name: str,
    media_type: str,
) -> NormalizedDocument:
    document = Document(str(path))
    segments = [
        NormalizedSegment(locator={"paragraph": index}, text=paragraph.text)
        for index, paragraph in enumerate(document.paragraphs, 1)
        if paragraph.text
    ]
    return NormalizedDocument(
        object_sha256=digest,
        media_type=media_type,
        kind="docx",
        source_name=source_name,
        segments=segments,
        metadata={"paragraphs": len(document.paragraphs)},
    )
